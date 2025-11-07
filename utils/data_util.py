import sqlite3
import logging
from typing import Optional, List, Dict, Any
from pathlib import Path
import streamlit as st
import os
from docx import Document
from datetime import datetime
from difflib import SequenceMatcher
import pandas as pd
import random


def get_db_connection() -> Optional[sqlite3.Connection]:
    """获取数据库连接（连接到现有SQLite数据库）"""
    try:
        # 优先使用环境变量中的数据目录，如果没有则使用默认路径
        if os.environ.get("DOCUMANAGER_DATA_DIR"):
            base_dir = Path(os.environ["DOCUMANAGER_DATA_DIR"])
        else:
            base_dir = Path("./database")
        
        # 数据库文件路径
        db_path = base_dir / "sqlite_database.db"
        
        # 如果数据库文件不存在，返回None
        if not db_path.exists():
            logging.warning(f"数据库文件不存在: {db_path}")
            return None
        
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row  # 支持按字段名访问
        return conn
    except sqlite3.Error as e:
        logging.error(f"数据库连接失败: {str(e)}")
        return None

def get_db_statistics():
    conn = get_db_connection()
    
    result = {
        "document_count": 0,
    }
    try:
        if conn is None:
            raise RuntimeError("无法连接到数据库")
        cursor = conn.cursor()

        try:
            cursor.execute("""
                SELECT COUNT(*) AS total
                FROM document
                WHERE 1=1
            """)
            col_data = cursor.fetchone()
            result["document_count"] = col_data[0]
        except sqlite3.Error as e:
            raise RuntimeError(f"查询失败：{str(e)}")
    except sqlite3.Error as e:
        raise RuntimeError(f"数据库连接失败：{str(e)}")
    finally:
        if conn:
            conn.close()
    
    return result


def match_files_by_name(doc_folder: str = "./data/documents", img_folder: str = "./data/images") -> Dict:
    """
    对比documents和images文件夹内的文件名称，首先进行完全匹配
    
    参数：
        doc_folder: documents文件夹路径
        img_folder: images文件夹路径
    
    返回：
        dict: 包含匹配结果的字典
            {
                "status": "success" 或 "error",
                "matched_files": 完全匹配的文件对列表,
                "unmatched_docs": 未匹配的文档文件列表,
                "unmatched_images": 未匹配的图片文件列表,
                "error_msg": 错误信息（仅status为error时存在）
            }
    """
    result = {
        "status": "success",
        "matched_files": [],
        "unmatched_docs": [],
        "unmatched_images": []
    }
    
    try:
        # 获取文档文件列表
        doc_files = []
        if os.path.exists(doc_folder) and os.path.isdir(doc_folder):
            for file in os.listdir(doc_folder):
                if os.path.isfile(os.path.join(doc_folder, file)) and not file.startswith('.'):
                    # 获取文件名主体（不含扩展名）
                    file_main = file.rsplit('.', 1)[0] if '.' in file else file
                    doc_files.append({
                        "filename": file,
                        "name": file_main,
                        "absolute_path": os.path.abspath(os.path.join(doc_folder, file))
                    })
        
        # 获取图片文件列表
        img_files = []
        if os.path.exists(img_folder) and os.path.isdir(img_folder):
            for file in os.listdir(img_folder):
                if os.path.isfile(os.path.join(img_folder, file)) and not file.startswith('.'):
                    # 获取文件名主体（不含扩展名）
                    file_main = file.rsplit('.', 1)[0] if '.' in file else file
                    img_files.append({
                        "filename": file,
                        "name": file_main,
                        "absolute_path": os.path.abspath(os.path.join(img_folder, file))
                    })
        
        # 建立图片文件名到文件信息的映射
        img_name_map = {img["name"]: img for img in img_files}
        
        # 进行完全匹配
        matched_doc_names = set()
        matched_img_names = set()
        
        for doc in doc_files:
            if doc["name"] in img_name_map:
                # 找到完全匹配的图片
                img = img_name_map[doc["name"]]
                result["matched_files"].append({
                    "document": doc,
                    "image": img,
                    "match_type": "exact"
                })
                matched_doc_names.add(doc["name"])
                matched_img_names.add(doc["name"])
            else:
                # 未匹配的文档
                result["unmatched_docs"].append(doc)
        
        # 找出未匹配的图片
        for img in img_files:
            if img["name"] not in matched_img_names:
                result["unmatched_images"].append(img)
        
    except Exception as e:
        result["status"] = "error"
        result["error_msg"] = f"文件匹配失败：{str(e)}"
    
    return result

def match_files_with_similarity(doc_folder: str = "./data/documents", img_folder: str = "./data/images", threshold: float = 0.9) -> Dict:
    """
    对比documents和images文件夹内的文件名称，先完全匹配，再使用相似度匹配
    
    参数：
        doc_folder: documents文件夹路径
        img_folder: images文件夹路径
        threshold: 相似度匹配阈值
    
    返回：
        dict: 包含所有匹配结果的字典
    """
    # 首先进行完全匹配
    exact_match_result = match_files_by_name(doc_folder, img_folder)
    
    if exact_match_result["status"] == "error":
        return exact_match_result
    
    # 获取未匹配的文档和图片
    unmatched_docs = exact_match_result["unmatched_docs"]
    unmatched_images = exact_match_result["unmatched_images"]
    
    # 相似度匹配
    similarity_matches = []
    matched_img_indices = set()
    
    for doc_idx, doc in enumerate(unmatched_docs):
        best_match = None
        best_similarity = 0
        best_img_idx = -1
        
        for img_idx, img in enumerate(unmatched_images):
            if img_idx in matched_img_indices:
                continue
                
            # 计算文件名相似度
            similarity = SequenceMatcher(None, doc["name"], img["name"]).ratio()
            
            if similarity > best_similarity:
                best_similarity = similarity
                best_match = img
                best_img_idx = img_idx
        
        # 如果找到相似度高于阈值的匹配
        if best_match and best_similarity >= threshold:
            similarity_matches.append({
                "document": doc,
                "image": best_match,
                "match_type": "similarity",
                "similarity": best_similarity
            })
            matched_img_indices.add(best_img_idx)
    
    # 更新结果
    result = exact_match_result.copy()
    result["matched_files"].extend(similarity_matches)
    
    # 更新未匹配的文档和图片
    remaining_docs = []
    for doc_idx, doc in enumerate(unmatched_docs):
        # 检查该文档是否通过相似度匹配到了图片
        matched = False
        for match in similarity_matches:
            if match["document"]["absolute_path"] == doc["absolute_path"]:
                matched = True
                break
        if not matched:
            remaining_docs.append(doc)
    
    remaining_images = []
    for img_idx, img in enumerate(unmatched_images):
        if img_idx not in matched_img_indices:
            remaining_images.append(img)
    
    result["unmatched_docs"] = remaining_docs
    result["unmatched_images"] = remaining_images
    
    return result

def batch_insert_matched_files(matched_files: List[Dict]) -> Dict:
    """
    批量将匹配的文件对写入数据库
    
    参数：
        matched_files: 匹配的文件对列表
    
    返回：
        dict: 包含执行结果的字典
    """
    result = {
        "status": "success",
        "inserted_count": 0,
        "failed_count": 0,
        "error_msg": ""
    }
    
    conn = None  # 确保conn总是被初始化
    try:
        conn = get_db_connection()
        if conn is None:
            raise RuntimeError("数据库连接失败，无法获取游标")
        cursor = conn.cursor()
        
        for match in matched_files:
            try:
                doc = match["document"]
                img = match["image"]
                
                # 获取文档信息
                file_name = doc["name"]
                # 从name中解析作者、日期和文档名
                try:
                    author_name, publishdate, document_name = doc["name"].split("-")
                except ValueError:
                    author_name = "未知"
                    publishdate = ""
                    document_name = doc["name"]
                
                # 获取文件内容
                content = ""
                if os.path.exists(doc["absolute_path"]) and doc["absolute_path"].endswith(('.docx', '.doc')):
                    doc_obj = Document(doc["absolute_path"])
                    content = '\n'.join([para.text.strip() for para in doc_obj.paragraphs if para.text.strip()])
                
                create_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                
                # 插入数据，包含图片路径信息
                cursor.execute(
                    "INSERT INTO document (filename, mediafilename, documentname, authorname, publishdate, created_at, content) VALUES (?, ?, ?, ?, ?, ?, ?)",
                    (file_name, img["filename"], document_name, author_name, publishdate, create_time, content)
                )
                result["inserted_count"] += 1
                
                logging.info(f"成功插入文档：{file_name}")

            except Exception as e:
                result["failed_count"] += 1
        
        conn.commit()
        conn.close()
        
    except Exception as e:
        result["status"] = "error"
        result["error_msg"] = f"批量插入失败：{str(e)}"
    finally:
        if conn:
            conn.close()
    
    return result

def insert_single_file(doc_file, img_file, author_name=None, publishdate=None) -> Dict:
    """
    将单个匹配的文件对写入数据库并保存文件到本地目录
    
    参数：
        doc_file: Streamlit上传的文件对象
        img_file: Streamlit上传的文件对象
        author_name: 作者名称（可选）
        publishdate: 发布日期（可选）
    
    返回：
        dict: 包含执行结果的字典
    """
    result = {
        "status": "success",
        "inserted_count": 0,
        "error_msg": ""
    }
    
    conn = None
    doc_path = None
    img_path = None
    try:
        # 获取文件名
        doc_filename = doc_file.name
        # 确保img_filename始终是字符串，即使img_file为空
        img_filename = img_file.name if img_file else ""
        doc_name = doc_filename.rsplit('.', 1)[0] if '.' in doc_filename else doc_filename
        
        # 从环境变量获取数据目录，如果没有则使用默认值
        data_dir = os.environ.get('DOCUMANAGER_DATA_DIR', './data')
        
        # 创建数据目录
        docs_dir = Path(data_dir) / "documents"
        images_dir = Path(data_dir) / "images"
        docs_dir.mkdir(parents=True, exist_ok=True)
        images_dir.mkdir(parents=True, exist_ok=True)
        
        # 保存文件到指定目录
        doc_path = docs_dir / doc_filename
        
        # 写入文档文件
        with open(doc_path, "wb") as f:
            f.write(doc_file.getvalue())
        
        # 写入图片文件
        img_path = None
        if img_file and img_filename:
            # 确保Path操作时img_filename是有效的字符串
            img_path = images_dir / str(img_filename)
            with open(img_path, "wb") as f:
                f.write(img_file.getvalue())
        
        # 连接数据库
        conn = get_db_connection()
        if conn is None:
            raise RuntimeError("数据库连接失败，无法获取游标")
        cursor = conn.cursor()
        
        # 文档名称和信息解析
        document_name = doc_name
        
        # 如果未提供作者信息，尝试从文件名解析
        if author_name is None:
            try:
                if '-' in doc_name:
                    parts = doc_name.split("-")
                    if len(parts) >= 3:
                        author_name = parts[0]
                        publishdate = parts[1] if publishdate is None else publishdate
                        document_name = "-".join(parts[2:])
                    else:
                        author_name = "未知"
                else:
                    author_name = "未知"
            except:
                author_name = "未知"
        
        # 读取文档内容（仅支持docx/doc）
        content = ""
        if doc_filename.endswith(('.docx', '.doc')):
            try:
                doc_obj = Document(doc_path)
                content = '\n'.join([para.text.strip() for para in doc_obj.paragraphs if para.text.strip()])
            except Exception as e:
                logging.warning(f"读取文档内容失败：{str(e)}")
                content = ""
        
        create_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 插入数据
        cursor.execute(
            "INSERT INTO document (filename, mediafilename, documentname, authorname, publishdate, created_at, content) VALUES (?, ?, ?, ?, ?, ?, ?)",
            (doc_filename, img_filename, document_name, author_name, publishdate, create_time, content)
        )
        result["inserted_count"] = 1
        
        conn.commit()
        logging.info(f"成功插入单个文档：{doc_filename}，图片：{img_filename}")
        
    except Exception as e:
        result["status"] = "error"
        result["error_msg"] = f"单文件插入失败：{str(e)}"
        # 发生错误时尝试删除已保存的文件
        try:
            if doc_path is not None and doc_path.exists():
                doc_path.unlink()
            if img_path is not None and img_path.exists():
                img_path.unlink()
        except:
            pass
    finally:
        if conn:
            conn.close()
    
    return result

def search_records(filters: Dict[str, Any]) -> pd.DataFrame:
    conn = get_db_connection()
    if conn is None:
        return pd.DataFrame()
    
    try:
        # 构建基础查询
        query = "SELECT * FROM document WHERE 1=1"
        params = []
        
        # 添加过滤条件
        if filters.get('filename'):
            query += " AND filename LIKE ?"
            params.append(f"%{filters['filename']}%")
        
        if filters.get('mediafilename'):
            query += " AND mediafilename LIKE ?"
            params.append(f"%{filters['mediafilename']}%")
        
        if filters.get('documentname'):
            query += " AND documentname LIKE ?"
            params.append(f"%{filters['documentname']}%")
        
        if filters.get('authorname'):
            query += " AND authorname LIKE ?"
            params.append(f"%{filters['authorname']}%")
        
        if filters.get('full_text'):
            query += " AND content LIKE ?"
            params.append(f"%{filters['full_text']}%")
        
        # 日期范围筛选（新增）
        if filters.get('start_date') and filters.get('end_date'):
            query += " AND publishdate BETWEEN ? AND ?"
            params.extend([filters['start_date'], filters['end_date']])
        elif filters.get('start_date'):
            query += " AND publishdate >= ?"
            params.append(filters['start_date'])
        elif filters.get('end_date'):
            query += " AND publishdate <= ?"
            params.append(filters['end_date'])
        
        # 执行查询
        df = pd.read_sql_query(query, conn, params=params)
        # 转换日期格式为字符串
        if 'publishdate' in df.columns:
            df['publishdate'] = pd.to_datetime(df['publishdate'], errors='coerce').dt.strftime('%Y-%m-%d')
        if 'created_at' in df.columns:
            df['created_at'] = pd.to_datetime(df['created_at'], errors='coerce').dt.strftime('%Y-%m-%d %H:%M:%S')
        return df
    except Exception as e:
        st.error(f"查询失败：{str(e)}")
        return pd.DataFrame()
    finally:
        conn.close()

# 获取随机一条记录
def get_random_record() -> Optional[Dict[str, Any]]:
    conn = get_db_connection()
    if conn is None:
        return None
    
    try:
        cursor = conn.cursor()
        # 先获取记录总数
        cursor.execute("SELECT COUNT(*) FROM document")
        count = cursor.fetchone()[0]
        
        if count == 0:
            st.info("数据库中暂无记录")
            return None
        
        # 随机选择一条记录
        random_offset = random.randint(0, count - 1)
        cursor.execute("SELECT * FROM document LIMIT 1 OFFSET ?", (random_offset,))
        record = cursor.fetchone()
        # 转换为字典并格式化日期
        record_dict = {k: record[k] for k in record.keys()}
        if 'publishdate' in record_dict:
            record_dict['publishdate'] = pd.to_datetime(record_dict['publishdate'], errors='coerce').strftime('%Y-%m-%d')
        if 'created_at' in record_dict:
            record_dict['created_at'] = pd.to_datetime(record_dict['created_at'], errors='coerce').strftime('%Y-%m-%d %H:%M:%S')
        return record_dict
    except Exception as e:
        st.error(f"获取随机记录失败：{str(e)}")
        return None

# 分页处理函数
def get_paginated_data(df: pd.DataFrame, page: int, page_size: int) -> pd.DataFrame:
    start_idx = (page - 1) * page_size
    end_idx = start_idx + page_size
    return df.iloc[start_idx:end_idx].copy()